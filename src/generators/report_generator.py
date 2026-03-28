"""Generate monthly analysis report as Word document"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from decimal import Decimal
from datetime import datetime
from typing import Dict, Any
from src.models import RentRollData, T12FinancialData


class ReportGenerator:
    """Generate monthly analysis reports in Word format"""

    def __init__(self, rent_roll: RentRollData, t12_data: T12FinancialData,
                 occupancy_analysis: Dict[str, Any], financial_analysis: Dict[str, Any],
                 collections_analysis: Dict[str, Any]):
        self.rent_roll = rent_roll
        self.t12_data = t12_data
        self.occupancy_analysis = occupancy_analysis
        self.financial_analysis = financial_analysis
        self.collections_analysis = collections_analysis

    def generate(self, output_path: str) -> str:
        """Generate report and save to path"""
        doc = Document()

        # Add cover page
        self._add_cover_page(doc)

        # Add sections
        self._add_occupancy_section(doc)
        self._add_rental_section(doc)
        self._add_financial_section(doc)
        self._add_collections_section(doc)
        self._add_appendices(doc)

        # Save
        doc.save(output_path)
        return output_path

    def _add_cover_page(self, doc):
        """Add cover page"""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run('Monthly Analysis Report')
        title_run.font.size = Pt(24)
        title_run.font.bold = True

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run(self.rent_roll.property_name)
        subtitle_run.font.size = Pt(18)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Report Date: {self.rent_roll.report_date.strftime('%B %d, %Y')}")
        date_run.font.size = Pt(12)

        doc.add_paragraph()  # Spacer

    def _add_occupancy_section(self, doc):
        """Add occupancy analysis section"""
        doc.add_heading('A. Occupancy Analysis', level=1)

        # Summary metrics
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        metrics = [
            ('Total Units', str(self.occupancy_analysis['total_units'])),
            ('Occupied Units', str(self.occupancy_analysis['occupied_units'])),
            ('Vacant Units', str(self.occupancy_analysis['vacant_units'])),
            ('Occupancy Rate', f"{self.occupancy_analysis['physical_occupancy_pct']}%"),
            ('Vacancy Loss', f"${self.occupancy_analysis['vacancy_loss']:,.2f}"),
        ]

        for metric, value in metrics:
            row = table.add_row()
            row_cells = row.cells
            row_cells[0].text = metric
            row_cells[1].text = value

        # Vacant units list
        if self.occupancy_analysis['vacant_units_detail']:
            doc.add_paragraph()
            doc.add_paragraph(f"Vacant Units: {', '.join(self.occupancy_analysis['vacant_units_detail'])}")

    def _add_rental_section(self, doc):
        """Add rental analysis section"""
        doc.add_heading('B. Rental Analysis', level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Amount'

        metrics = [
            ('Market Rent (100% Occupied)', f"${self.rent_roll.market_rent:,.2f}"),
            ('Actual Rent Collected', f"${self.rent_roll.actual_rent:,.2f}"),
            ('Loss to Lease', f"${self.rent_roll.loss_to_lease:,.2f}"),
            ('Vacancy Loss', f"${self.rent_roll.vacancy_loss:,.2f}"),
        ]

        for metric, value in metrics:
            row = table.add_row()
            row_cells = row.cells
            row_cells[0].text = metric
            row_cells[1].text = value

    def _add_financial_section(self, doc):
        """Add financial analysis section"""
        doc.add_heading('C. Financial Analysis', level=1)

        # Revenue trends
        doc.add_heading('C.1 Revenue Analysis', level=2)
        revenue = self.financial_analysis['revenue_analysis']

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Month'
        header_cells[1].text = 'Rental Income'
        header_cells[2].text = 'Other Income'
        header_cells[3].text = 'Total Income'
        header_cells[4].text = 'Trend'

        for monthly in revenue['monthly_totals'][-3:]:  # Last 3 months
            row = table.add_row()
            cells = row.cells
            cells[0].text = monthly['month']
            cells[1].text = f"${monthly['rental']:,.2f}"
            cells[2].text = f"${monthly['other']:,.2f}"
            cells[3].text = f"${monthly['total']:,.2f}"
            cells[4].text = revenue['trend']

        # Expense analysis
        doc.add_heading('C.2 Expense Analysis', level=2)
        expenses = self.financial_analysis['expense_analysis']

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Category'
        header_cells[1].text = 'Average Monthly'

        for category, avg in expenses['category_averages'].items():
            row = table.add_row()
            cells = row.cells
            cells[0].text = category.replace('_', ' ').title()
            cells[1].text = f"${avg:,.2f}"

        # NOI
        doc.add_heading('C.3 Net Operating Income', level=2)
        noi = self.financial_analysis['noi_analysis']
        para = doc.add_paragraph(f"Average NOI: ${noi['average_noi']:,.2f}")

    def _add_collections_section(self, doc):
        """Add collections analysis section"""
        doc.add_heading('D. Collections Analysis', level=1)

        collections = self.collections_analysis
        stats = collections['collection_stats']

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'

        metrics = [
            ('Total Charged', f"${stats['total_charged']:,.2f}"),
            ('Total Paid', f"${stats['total_paid']:,.2f}"),
            ('Collection Rate', f"{stats['collection_rate_pct']}%"),
            ('Delinquent Units', str(stats['total_delinquent_units'])),
            ('Total Delinquent Amount', f"${self.collections_analysis['total_debit_balance']:,.2f}"),
        ]

        for metric, value in metrics:
            row = table.add_row()
            row_cells = row.cells
            row_cells[0].text = metric
            row_cells[1].text = value

        # Top delinquent units
        if collections['delinquent_units']:
            doc.add_paragraph()
            doc.add_paragraph('Top Delinquent Units:')
            for unit in collections['delinquent_units'][:10]:
                doc.add_paragraph(
                    f"Unit {unit['unit_number']}: ${unit['debit_balance']:,.2f}",
                    style='List Bullet'
                )

    def _add_appendices(self, doc):
        """Add appendices with detailed data"""
        doc.add_page_break()
        doc.add_heading('Appendices', level=1)

        # Appendix 1: Unit details
        doc.add_heading('Appendix 1: Unit Details', level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Unit'
        header_cells[1].text = 'Tenant'
        header_cells[2].text = 'Status'
        header_cells[3].text = 'Market Rent'
        header_cells[4].text = 'Actual Rent'
        header_cells[5].text = 'Balance'

        for unit in self.rent_roll.units[:50]:  # Limit to first 50
            row = table.add_row()
            cells = row.cells
            cells[0].text = unit.unit_number
            cells[1].text = unit.tenant_name or '-'
            cells[2].text = unit.status
            cells[3].text = f"${unit.market_rent:,.2f}"
            cells[4].text = f"${unit.actual_rent:,.2f}"
            cells[5].text = f"${unit.debit_balance:,.2f}"

        if len(self.rent_roll.units) > 50:
            doc.add_paragraph(f"... and {len(self.rent_roll.units) - 50} more units")

    def _format_currency(self, value: Decimal) -> str:
        """Format Decimal as currency"""
        return f"${value:,.2f}"

    def _format_percentage(self, value: Decimal) -> str:
        """Format Decimal as percentage"""
        return f"{value:.2f}%"
