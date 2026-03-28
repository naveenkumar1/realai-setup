"""Generate report matching the template format exactly"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from decimal import Decimal
from datetime import datetime
from typing import Dict, Any
from src.models import RentRollData, T12FinancialData
from src.narrative_generator import NarrativeGenerator


class TemplateReportGenerator:
    """Generate report matching template format"""

    def __init__(self, rent_roll: RentRollData, t12_data: T12FinancialData,
                 occupancy_analysis: Dict[str, Any], financial_analysis: Dict[str, Any],
                 collections_analysis: Dict[str, Any]):
        self.rent_roll = rent_roll
        self.t12_data = t12_data
        self.occupancy_analysis = occupancy_analysis
        self.financial_analysis = financial_analysis
        self.collections_analysis = collections_analysis
        self.narrative_gen = NarrativeGenerator()

    def generate(self, output_path: str) -> str:
        """Generate report matching template format"""
        doc = Document()

        # Set margins to match template
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Cover page
        self._add_cover_page(doc)

        # Audit top sheet (summary table)
        self._add_audit_top_sheet(doc)

        # Operational report analysis
        self._add_operational_section(doc)

        # Financial analysis
        self._add_financial_section(doc)

        # Appendices
        self._add_appendices(doc)

        doc.save(output_path)
        return output_path

    def _add_cover_page(self, doc):
        """Add cover page matching template"""
        # Header
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("GW2R Asset Management, LLC").font.size = Pt(10)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run("Frisco, Texas, USA").font.size = Pt(10)

        doc.add_paragraph()
        doc.add_paragraph()

        # Title
        p = doc.add_paragraph("Asset Management Audit Report")
        p_run = p.runs[0]
        p_run.font.bold = True
        p_run.font.size = Pt(14)

        # Property name and date
        p = doc.add_paragraph(self.rent_roll.property_name)
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.bold = True

        p = doc.add_paragraph(self.rent_roll.report_date.strftime("%B %Y"))
        p.runs[0].font.size = Pt(11)

        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph("Audit Top Sheet")
        p.runs[0].font.bold = True
        p.runs[0].font.size = Pt(11)

    def _add_audit_top_sheet(self, doc):
        """Add audit top sheet with key findings"""
        doc.add_paragraph()
        doc.add_paragraph()

        # Create audit findings table
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        headers = ['Report Section', 'Audit Findings', 'Actionable Items', 'Reference', 'Deadline']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Add operational findings
        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Operational Report Analysis'
        cells[1].text = f"Occupancy: {self.occupancy_analysis['physical_occupancy_pct']}%. Vacant units: {self.occupancy_analysis['vacant_units']}"
        cells[2].text = "Strengthen leasing efforts through targeted marketing"
        cells[3].text = ""
        cells[4].text = ""

        # Add financial findings
        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Financial Analysis'
        cells[1].text = f"Average Revenue: ${self.financial_analysis['revenue_analysis']['average_monthly_revenue']:,.0f}. Delinquent: ${self.collections_analysis['total_debit_balance']:,.0f}"
        cells[2].text = "Implement stronger collection procedures"
        cells[3].text = ""
        cells[4].text = ""

    def _add_operational_section(self, doc):
        """Add operational report analysis section"""
        doc.add_page_break()

        heading = doc.add_heading('OPERATIONAL REPORT ANALYSIS', level=1)
        heading.runs[0].font.size = Pt(12)

        # A.1 Occupancy Analysis
        doc.add_heading('A.1 Occupancy Analysis:', level=2)
        doc.add_paragraph()

        # Generate narrative
        occupancy_narrative = self.narrative_gen.generate_occupancy_narrative(
            self.occupancy_analysis, self.rent_roll.__dict__
        )
        doc.add_paragraph(occupancy_narrative)

        # Add occupancy metrics
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        months = ['OCCUPANCY ANALYSIS', 'Latest Month', 'Prior Month', 'Prior Month -1', 'Prior Month -2', 'Prior Month -3', 'Prior Month -4']
        for i, month in enumerate(months):
            header_cells[i].text = month

        # Add data row
        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Total units'
        cells[1].text = str(self.occupancy_analysis['total_units'])

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Physical Occupancy'
        cells[1].text = f"{self.occupancy_analysis['physical_occupancy_pct']}%"

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Vacancy'
        cells[1].text = str(self.occupancy_analysis['vacant_units'])

        # A.2 Rental Analysis
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_heading('A.2 Rental Analysis:', level=2)
        doc.add_paragraph()

        rental_narrative = self.narrative_gen.generate_rental_narrative({
            'market_rent': self.rent_roll.market_rent,
            'actual_rent': self.rent_roll.actual_rent,
            'loss_to_lease': self.rent_roll.loss_to_lease,
            'vacancy_loss': self.rent_roll.vacancy_loss
        })
        doc.add_paragraph(rental_narrative)

        # Rental metrics table
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'RENTAL ANALYSIS'
        header_cells[1].text = 'Latest Month'
        header_cells[2].text = 'Prior Month'
        header_cells[3].text = 'Prior Month -1'

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Market Rent'
        cells[1].text = f"${self.rent_roll.market_rent:,.2f}"

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Actual Rent'
        cells[1].text = f"${self.rent_roll.actual_rent:,.2f}"

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Loss to Lease'
        cells[1].text = f"${self.rent_roll.loss_to_lease:,.2f}"

    def _add_financial_section(self, doc):
        """Add financial analysis section"""
        doc.add_paragraph()
        doc.add_page_break()

        heading = doc.add_heading('FINANCIAL ANALYSIS', level=1)
        heading.runs[0].font.size = Pt(12)

        # C.1 Profit and Loss
        doc.add_heading('C.1 Profit and Loss Analysis:', level=2)
        doc.add_paragraph()

        financial_narrative = self.narrative_gen.generate_financial_narrative(
            self.financial_analysis, self.rent_roll.report_date.strftime("%B %Y")
        )
        doc.add_paragraph(financial_narrative)

        # P&L Table
        doc.add_paragraph()
        revenue = self.financial_analysis['revenue_analysis']
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Latest Month'
        header_cells[2].text = 'Prior Month'
        header_cells[3].text = 'Prior Month -1'
        header_cells[4].text = 'Trend'

        if revenue['monthly_totals']:
            latest = revenue['monthly_totals'][-1]
            row = table.add_row()
            cells = row.cells
            cells[0].text = 'Total Income'
            cells[1].text = f"${latest['total']:,.2f}"
            cells[4].text = revenue['trend']

            row = table.add_row()
            cells = row.cells
            cells[0].text = 'Total Expenses'
            cells[1].text = f"${self.financial_analysis['expense_analysis']['total_annual_expenses']/12:,.2f}"

            row = table.add_row()
            cells = row.cells
            cells[0].text = 'NOI'
            cells[1].text = f"${self.financial_analysis['noi_analysis']['average_noi']:,.2f}"

        # C.3 Collection Analysis
        doc.add_paragraph()
        doc.add_heading('C.3 Collection Analysis:', level=2)
        doc.add_paragraph()

        collections_narrative = self.narrative_gen.generate_collections_narrative(
            self.collections_analysis
        )
        doc.add_paragraph(collections_narrative)

        # Collections table
        doc.add_paragraph()
        stats = self.collections_analysis['collection_stats']
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Amount'

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Total Charged'
        cells[1].text = f"${stats['total_charged']:,.2f}"

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Total Paid'
        cells[1].text = f"${stats['total_paid']:,.2f}"

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Collection Rate'
        cells[1].text = f"{stats['collection_rate_pct']}%"

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Delinquent Units'
        cells[1].text = str(stats['total_delinquent_units'])

        row = table.add_row()
        cells = row.cells
        cells[0].text = 'Total Delinquent Balance'
        cells[1].text = f"${self.collections_analysis['total_debit_balance']:,.2f}"

    def _add_appendices(self, doc):
        """Add appendices"""
        doc.add_page_break()

        heading = doc.add_heading('APPENDICES', level=1)
        heading.runs[0].font.size = Pt(12)

        # Appendix with unit details
        doc.add_heading('Unit Details Report', level=2)

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        headers = ['Unit', 'Tenant', 'Status', 'Market Rent', 'Actual Rent', 'Balance']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        # Add unit rows (limit to 30 for readability)
        for unit in self.rent_roll.units[:30]:
            row = table.add_row()
            cells = row.cells
            cells[0].text = unit.unit_number
            cells[1].text = unit.tenant_name or '-'
            cells[2].text = unit.status
            cells[3].text = f"${unit.market_rent:,.2f}"
            cells[4].text = f"${unit.actual_rent:,.2f}"
            cells[5].text = f"${unit.debit_balance:,.2f}"

        if len(self.rent_roll.units) > 30:
            doc.add_paragraph(f"... and {len(self.rent_roll.units) - 30} more units")

        # Summary footer
        doc.add_page_break()
        doc.add_paragraph()
        doc.add_paragraph()

        p = doc.add_paragraph("Prepared by:")
        p.runs[0].font.bold = True

        p = doc.add_paragraph("                            Checked by:                                       Verified by:")
        p.runs[0].font.size = Pt(10)

        p = doc.add_paragraph("____________________          _____________________                  ___________________")
        p.runs[0].font.size = Pt(10)

        p = doc.add_paragraph("Asset Manager                 Financial Controller/Auditor           President & CEO")
        p.runs[0].font.size = Pt(10)

        p = doc.add_paragraph("GW2R LLC                       GW2R LLC                              GW2R LLC")
        p.runs[0].font.size = Pt(10)
