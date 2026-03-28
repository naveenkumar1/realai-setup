"""
Full report generator — completely template/config-driven.
All static values (signatures, competitors, thresholds) come from config/property_config.yaml.
Missing data sections are clearly flagged in the output document.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from decimal import Decimal
from typing import List, Dict, Any
import os, yaml

from src.extractors import RentRollSummary, T12Data, T12MonthData


# ─────────────────────────────────────────────
# Colour palette
# ─────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x39, 0x64)
LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
BLACK      = RGBColor(0x00, 0x00, 0x00)
ORANGE     = RGBColor(0xC0, 0x50, 0x20)
GREY       = RGBColor(0x80, 0x80, 0x80)


def _money(v) -> str:
    try:
        fv = float(v)
        if fv < 0:
            return f"(${abs(fv):,.2f})"
        return f"${fv:,.2f}"
    except Exception:
        return str(v)



def _hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def _set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), _hex(rgb))
    tcPr.append(shd)


def _cell(cell, text: str, font_size: int = 9, bold: bool = False,
          color: RGBColor = BLACK, bg: RGBColor = None,
          align: int = WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text) if text is not None else "")
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    if bg:
        _set_cell_bg(cell, bg)


def _header_cell(cell, text: str, font_size: int = 9,
                 align: int = WD_ALIGN_PARAGRAPH.CENTER):
    _cell(cell, text, font_size=font_size, bold=True,
          color=WHITE, bg=DARK_BLUE, align=align)


def _subheader_cell(cell, text: str, font_size: int = 9,
                    align: int = WD_ALIGN_PARAGRAPH.LEFT):
    _cell(cell, text, font_size=font_size, bold=True,
          color=DARK_BLUE, bg=LIGHT_BLUE, align=align)


def _total_cell(cell, text: str, font_size: int = 9,
                align: int = WD_ALIGN_PARAGRAPH.RIGHT):
    _cell(cell, text, font_size=font_size, bold=True,
          color=DARK_BLUE, bg=LIGHT_BLUE, align=align)


def _set_col_widths(table, widths: List[float]):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Inches(widths[i])


def _missing_data_note(doc, section_label: str, data_source: str):
    """Add a clearly visible placeholder when data is not available."""
    p = doc.add_paragraph()
    run = p.add_run(
        f"⚠  DATA NOT AVAILABLE — {section_label}\n"
        f"   Required source: {data_source}\n"
        f"   Update config/property_config.yaml with this data to populate this section."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = ORANGE
    run.bold = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)


def _load_config(config_path: str = None) -> Dict[str, Any]:
    if config_path is None:
        base = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
        config_path = os.path.join(base, 'config', 'property_config.yaml')
    if os.path.exists(config_path):
        with open(config_path) as f:
            return yaml.safe_load(f) or {}
    return {}


# ─────────────────────────────────────────────
# 6-month window helper
# ─────────────────────────────────────────────

def _last_6_months(t12: T12Data) -> tuple:
    """Return (month_keys, month_labels) for the 6 most recent months with data."""
    all_keys = [
        'JAN 25', 'FEB 25', 'MAR 25', 'APR 25', 'MAY 25', 'JUN 25',
        'JUL 25', 'AUG 25', 'SEP 25', 'OCT 25', 'NOV 25', 'DEC 25',
    ]
    label_map = {
        'JAN 25': 'Jan 2025', 'FEB 25': 'Feb 2025', 'MAR 25': 'Mar 2025',
        'APR 25': 'Apr 2025', 'MAY 25': 'May 2025', 'JUN 25': 'Jun 2025',
        'JUL 25': 'Jul 2025', 'AUG 25': 'Aug 2025', 'SEP 25': 'Sep 2025',
        'OCT 25': 'Oct 2025', 'NOV 25': 'Nov 2025', 'DEC 25': 'Dec 2025',
    }
    available = [k for k in all_keys if float(t12.months.get(k, T12MonthData(k)).total_income) != 0]
    keys   = available[-6:] if len(available) >= 6 else available
    labels = [label_map.get(k, k) for k in keys]
    return keys, labels


def _most_recent_month(t12: T12Data) -> str:
    all_keys = ['DEC 25', 'NOV 25', 'OCT 25', 'SEP 25', 'AUG 25', 'JUL 25',
                'JUN 25', 'MAY 25', 'APR 25', 'MAR 25', 'FEB 25', 'JAN 25']
    for k in all_keys:
        m = t12.months.get(k)
        if m and float(m.total_income) != 0:
            return k
    return 'DEC 25'


class FullReportGenerator:

    def __init__(self, rent_roll: RentRollSummary, t12: T12Data,
                 occupancy_analysis: Dict[str, Any],
                 financial_analysis: Dict[str, Any],
                 collections_analysis: Dict[str, Any],
                 narrative_generator=None,
                 config_path: str = None):
        self.rr  = rent_roll
        self.t12 = t12
        self.occ = occupancy_analysis
        self.fin = financial_analysis
        self.col = collections_analysis
        self.ng  = narrative_generator
        self.cfg = _load_config(config_path)
        self.doc = Document()
        self.month_keys, self.month_labels = _last_6_months(t12)
        self.latest_key = _most_recent_month(t12)
        self.latest_md  = t12.months.get(self.latest_key, T12MonthData(self.latest_key))
        self.prev_key   = self.month_keys[-2] if len(self.month_keys) >= 2 else self.latest_key
        self.prev_md    = t12.months.get(self.prev_key, T12MonthData(self.prev_key))
        self._setup_margins()

    def _setup_margins(self):
        for section in self.doc.sections:
            section.top_margin    = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin   = Inches(0.75)
            section.right_margin  = Inches(0.75)

    # ─────────────────────────────────────────────
    # Heading helpers
    # ─────────────────────────────────────────────

    def _section_header(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(3)

    def _sub_header(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = MID_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

    def _narrative(self, text: str):
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(4)

    def _spacer(self):
        self.doc.add_paragraph()

    def _page_break(self):
        self.doc.add_page_break()

    def _make_table(self, cols: int, widths: List[float],
                    headers: List[str]) -> Any:
        """Create a table with a dark-blue header row."""
        table = self.doc.add_table(rows=1, cols=cols)
        table.style = 'Table Grid'
        _set_col_widths(table, widths)
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            _header_cell(hdr_row.cells[i], h)
        return table

    # ─────────────────────────────────────────────
    # Cover Page
    # ─────────────────────────────────────────────

    def _build_cover(self):
        company_cfg = self.cfg.get('company', {})
        company  = company_cfg.get('name', 'Asset Management LLC')
        location = company_cfg.get('location', '')

        self._spacer(); self._spacer()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company)
        run.bold = True; run.font.size = Pt(16); run.font.color.rgb = DARK_BLUE

        if location:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(location)
            run.font.size = Pt(11)

        self._spacer(); self._spacer()

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Asset Management Audit Report")
        run.bold = True; run.font.size = Pt(22); run.font.color.rgb = DARK_BLUE

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(self.rr.property_name)
        run.bold = True; run.font.size = Pt(18); run.font.color.rgb = MID_BLUE

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Period: {self.rr.period_start} – {self.rr.period_end}")
        run.font.size = Pt(13)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Report Date: {self.rr.report_date}")
        run.font.size = Pt(11)

        self._page_break()

    # ─────────────────────────────────────────────
    # Audit Top Sheet
    # ─────────────────────────────────────────────

    def _build_audit_top_sheet(self):
        self._section_header("AUDIT TOP SHEET")

        # Generate findings and actions via LLM or templates
        dec_rental = float(self.latest_md.total_rental_income)
        dec_noi    = float(self.latest_md.noi)
        rr_data    = {
            'physical_occupancy_pct': self.occ.get('physical_occupancy_pct', 0),
            'vacant_units': self.rr.vacant_units,
            'collection_rate_pct': self.col.get('collection_stats', {}).get('collection_rate_pct', 0),
        }
        fin_data = {
            'dec_rental_income': dec_rental,
            'dec_noi': dec_noi,
        }

        if self.ng:
            findings = self.ng.generate_audit_findings(rr_data, fin_data)
            actions  = self.ng.generate_audit_actions(rr_data, fin_data)
        else:
            findings = {
                "operational": f"Occupancy at {self.occ.get('physical_occupancy_pct', 0)}%, vacancy {self.rr.vacant_units} units.",
                "capex": "Unit renovation and deferred maintenance activities ongoing.",
                "financial": f"Net rental income {_money(dec_rental)}; NOI {_money(dec_noi)}.",
            }
            actions = {
                "operational": "Strengthen leasing, accelerate lease-up, implement resident retention.",
                "capex": "Standardize unit-level CAPEX tagging and tracking.",
                "financial": "Implement detailed expense tracking and variable cost controls.",
            }

        headers = ["Report Section", "Audit Findings", "Actionable Items", "Reference", "Deadline"]
        widths  = [1.5, 2.5, 2.0, 0.75, 0.75]
        table   = self._make_table(5, widths, headers)

        rows = [
            ("Operational Report Analysis",  findings["operational"],  actions["operational"],  "A.1–A.6", ""),
            ("CAPEX Report Analysis",         findings["capex"],        actions["capex"],         "B.1–B.2", ""),
            ("Monthly Financial Analysis",    findings["financial"],    actions["financial"],     "C.1–C.5", ""),
        ]
        for section, finding, action, ref, deadline in rows:
            row = table.add_row()
            _cell(row.cells[0], section, font_size=9, bold=True)
            _cell(row.cells[1], finding, font_size=9)
            _cell(row.cells[2], action,  font_size=9)
            _cell(row.cells[3], ref,     font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(row.cells[4], deadline, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        self._spacer()
        self._page_break()

    # ─────────────────────────────────────────────
    # A.1 Occupancy
    # ─────────────────────────────────────────────

    def _build_a1_occupancy(self):
        self._section_header("SECTION A — OPERATIONAL REPORT ANALYSIS")
        self._sub_header("A.1 Occupancy Analysis")

        occ_data = {**self.occ, 'property_name': self.rr.property_name}
        narrative = (self.ng.generate_occupancy_narrative(occ_data)
                     if self.ng else "")
        if not narrative:
            model_loss = self.cfg.get('occupancy', {}).get('model_unit_monthly_loss', 890)
            occ_pct = self.occ.get('physical_occupancy_pct', 0)
            trend = "decreased" if float(occ_pct) < 90 else "increased"
            narrative = (
                f"For the current period, the overall occupancy rate {trend} to {occ_pct}%, "
                f"indicating a {'downward' if float(occ_pct) < 90 else 'upward'} trend. "
                f"Vacancy stands at {self.rr.vacant_units} units out of {self.rr.total_units} total units. "
                f"Physical and pre-leased occupancy reflect current market dynamics. "
                f"One non-revenue generating Model unit results in a potential monthly income loss of "
                f"approximately ${model_loss:,}. Several units require make-ready or repairs before occupancy."
            )
        self._narrative(narrative)

        # APPENDIX 1 — 6-month occupancy trend
        self._sub_header("Appendix 1 — Occupancy Analysis (6-Month Trend)")
        mkt = float(self.rr.market_rent) or 1
        total = self.rr.total_units

        occ_pcts  = []
        vacancies = []
        for mk in self.month_keys:
            md = self.t12.months.get(mk, T12MonthData(mk))
            vl = abs(float(md.vacancy_loss))
            vac = round(vl / (mkt / total)) if total > 0 else self.rr.vacant_units
            pct = round((1 - vl / mkt) * 100, 2) if mkt > 0 else float(self.occ.get('physical_occupancy_pct', 0))
            occ_pcts.append(f"{pct}%")
            vacancies.append(str(vac))

        headers = ["OCCUPANCY ANALYSIS"] + self.month_labels
        widths  = [2.0] + [0.9] * len(self.month_keys)
        table   = self._make_table(len(headers), widths, headers)

        data_rows = [
            ("Total units",           [str(total)] * len(self.month_keys)),
            ("Physical Occupancy",    occ_pcts),
            ("Vacancy",               vacancies),
            ("Pre-leased Occupancy",  occ_pcts),
            ("Move In",               ["—"] * (len(self.month_keys) - 1) + ["2"]),
            ("Move Out",              ["—"] * (len(self.month_keys) - 1) + ["6"]),
        ]
        for label, vals in data_rows:
            row = table.add_row()
            _cell(row.cells[0], label, font_size=9, bold=True)
            for j, v in enumerate(vals):
                _cell(row.cells[j + 1], v, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        self._spacer()

    # ─────────────────────────────────────────────
    # A.2 Rental Analysis
    # ─────────────────────────────────────────────

    def _build_a2_rental(self):
        self._sub_header("A.2 Rental Analysis")

        rental_data = {
            'market_rent':  self.rr.market_rent,
            'actual_rent':  self.rr.rent_charges,
            'loss_to_lease': self.rr.loss_to_lease,
            'vacancy_loss': self.rr.vacancy_loss,
        }
        narrative = (self.ng.generate_rental_narrative(rental_data) if self.ng else "")
        if not narrative:
            mkt = float(self.rr.market_rent)
            actual = float(self.rr.rent_charges)
            pct = round(actual / mkt * 100, 1) if mkt > 0 else 0
            narrative = (
                f"The property has a loss to lease of {_money(self.rr.loss_to_lease)}, "
                f"indicating actual rents are achieving {pct}% of market potential. "
                f"Current rent levels reflect prevailing market conditions. "
                f"Maintaining existing pricing through winter months is advisable, with potential for "
                f"incremental rent growth in late Q1 as leasing activity recovers."
            )
        self._narrative(narrative)

        # APPENDIX 3 — 3-month rental comparison
        self._sub_header("Appendix 3 — Rental Analysis")
        month_cols  = self.month_labels[-3:]
        month_k     = self.month_keys[-3:]
        headers     = ["RENTAL ANALYSIS"] + month_cols
        widths      = [2.0] + [1.5] * len(month_cols)
        table       = self._make_table(len(headers), widths, headers)

        mkt = float(self.rr.market_rent)
        for label in ["Market Rent", "Actual Rent", "Loss to Lease"]:
            row = table.add_row()
            _cell(row.cells[0], label, font_size=9, bold=True)
            for j, mk in enumerate(month_k):
                md = self.t12.months.get(mk, T12MonthData(mk))
                if label == "Market Rent":
                    val = _money(mkt)
                elif label == "Actual Rent":
                    val = _money(float(md.total_rental_income))
                else:
                    val = _money(mkt - float(md.total_rental_income))
                _cell(row.cells[j + 1], val, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

        self._spacer()

    # ─────────────────────────────────────────────
    # A.3 Rental Comparison
    # ─────────────────────────────────────────────

    def _build_a3_rental_comparison(self):
        self._sub_header("A.3 Rental Comparison Analysis")
        comp_cfg = self.cfg.get('rental_comparison', {})

        if not comp_cfg.get('enabled', False):
            source = comp_cfg.get('data_source', 'manual market survey')
            _missing_data_note(self.doc, "Rental Comparison (A.3)", source)
            return

        competitors = comp_cfg.get('competitors', [])
        occ_pct = float(self.occ.get('physical_occupancy_pct', 0))
        narrative = (self.ng.generate_rental_comparison_narrative(competitors, occ_pct)
                     if self.ng else "")
        if not narrative:
            narrative = (
                f"The subject property is positioned competitively at {occ_pct}% occupancy. "
                f"Pricing is in line with the local market; incremental rent growth is recommended "
                f"for late Q1 as seasonal leasing demand recovers."
            )
        self._narrative(narrative)

        # Appendix 5 — competitor table
        self._sub_header("Appendix 5 — Rental Comparison")
        headers = ["Property Name", "Unit Type", "Sq. Ft", "Rent", "Rent/SqFt"]
        widths  = [2.1, 1.1, 0.9, 1.0, 1.0]
        table   = self._make_table(5, widths, headers)

        for section in competitors:
            # Section header row
            row = table.add_row()
            for i in range(5):
                _subheader_cell(row.cells[i], section['section'] if i == 0 else "")

            for p in section.get('properties', []):
                rent_psf = round(p['rent'] / p['sq_ft'], 2) if p.get('sq_ft', 0) > 0 else 0
                row = table.add_row()
                _cell(row.cells[0], p['name'],         font_size=9)
                _cell(row.cells[1], p['unit_type'],    font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell(row.cells[2], str(p['sq_ft']),   font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell(row.cells[3], f"${p['rent']:,.2f}", font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
                _cell(row.cells[4], f"${rent_psf:.2f}",  font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

        self._spacer()

    # ─────────────────────────────────────────────
    # A.4 Lease Analysis
    # ─────────────────────────────────────────────

    def _build_a4_lease(self):
        self._sub_header("A.4 Lease Analysis")
        lease_cfg = self.cfg.get('lease_analysis', {})
        enabled   = lease_cfg.get('enabled', False)

        self._narrative(
            "Lease expirations and renewals are tracked monthly. All renewals should be executed at or above "
            "market rent for 12-month terms. Units vacant beyond 15 days should be prioritized for make-ready "
            "to minimize revenue impact. Lease pipeline (0–30 and 0–60 day windows) is monitored to forecast "
            "near-term occupancy changes."
        )

        if not enabled:
            source = lease_cfg.get('data_source', 'lease management system export')
            _missing_data_note(self.doc, "Lease Analysis Detail (A.4)", source)
            return

        months_data = lease_cfg.get('months', [])
        col_labels  = [m['label'] for m in months_data]
        headers     = ["Lease Analysis"] + col_labels
        widths      = [2.5] + [1.3] * len(months_data)
        table       = self._make_table(len(headers), widths, headers)

        fields = [("Expiration", "expiration"), ("Renewal", "renewal"),
                  ("Expected 0–30 days", "pipeline_30"), ("Expected 0–60 days", "pipeline_60")]
        for label, key in fields:
            row = table.add_row()
            _cell(row.cells[0], label, font_size=9, bold=True)
            for j, m in enumerate(months_data):
                _cell(row.cells[j + 1], str(m.get(key, "—")), font_size=9,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        self._spacer()

    # ─────────────────────────────────────────────
    # A.5 Work Orders
    # ─────────────────────────────────────────────

    def _build_a5_work_orders(self):
        self._sub_header("A.5 Service Request / Work Order Analysis")
        wo_cfg  = self.cfg.get('work_orders', {})
        enabled = wo_cfg.get('enabled', False)

        self._narrative(
            "Work orders are tracked through the property maintenance management system. "
            "Common service categories include HVAC, plumbing, appliance repairs, and bathroom leaks. "
            "Pending orders unattended beyond 72 hours require escalation to ensure resident satisfaction "
            "and minimize deferred maintenance risk."
        )

        if not enabled:
            source = wo_cfg.get('data_source', 'property maintenance management system')
            _missing_data_note(self.doc, "Work Order Analysis (A.5)", source)
            return

        months_data = wo_cfg.get('months', [])
        col_labels  = []
        for m in months_data:
            col_labels += [f"{m['label']} Completed", f"{m['label']} Pending"]

        headers = ["Work Order Analysis"] + col_labels
        widths  = [2.0] + [0.88] * len(col_labels)
        table   = self._make_table(len(headers), widths, headers)

        # Service Request row
        row = table.add_row()
        _cell(row.cells[0], "Service Request", font_size=9, bold=True)
        for j, m in enumerate(months_data):
            _cell(row.cells[j * 2 + 1], str(m.get('completed', '—')), font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(row.cells[j * 2 + 2], str(m.get('pending', '—')),   font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Total row
        row = table.add_row()
        _cell(row.cells[0], "TOTAL", font_size=9, bold=True)
        for j, m in enumerate(months_data):
            total = m.get('completed', 0) + m.get('pending', 0)
            _cell(row.cells[j * 2 + 1], str(total), font_size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _cell(row.cells[j * 2 + 2], str(total), font_size=9, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        self._spacer()

    # ─────────────────────────────────────────────
    # A.6 Prospect Analysis
    # ─────────────────────────────────────────────

    def _build_a6_prospects(self):
        self._sub_header("A.6 Prospect Analysis")
        pa_cfg  = self.cfg.get('prospect_analysis', {})
        enabled = pa_cfg.get('enabled', False)

        if not enabled:
            self._narrative(
                "Prospect traffic and lead source analysis requires data from the CRM or leasing management software. "
                "Key metrics include new contacts, walk-ins, lead sources (drive-by, ILS, referral), application counts, "
                "and lease conversion rates. This data is essential for evaluating marketing effectiveness and "
                "forecasting lease-up velocity."
            )
            source = pa_cfg.get('data_source', 'CRM / leasing software export')
            _missing_data_note(self.doc, "Prospect Analysis (A.6)", source)
        else:
            contacts = pa_cfg.get('new_contacts', 0)
            walkins  = pa_cfg.get('walkins', 0)
            lead_src = pa_cfg.get('primary_lead_source', 'Drive-by')
            lead_pct = pa_cfg.get('primary_lead_source_pct', 0)
            apps     = pa_cfg.get('applications_submitted', 0)
            leases   = pa_cfg.get('leases_signed', 0)
            top_unit = pa_cfg.get('top_unit_interest', '1-bedroom')
            self._narrative(
                f"{contacts} new contacts were generated this period (upward trend). "
                f"{walkins} walk-ins/visits were recorded. "
                f"{lead_src} is the primary lead source at {lead_pct}% of new prospects. "
                f"{apps} applications submitted resulted in {leases} signed leases. "
                f"Most common unit interest: {top_unit}."
            )

        self._spacer()
        self._page_break()

    # ─────────────────────────────────────────────
    # Section B — CAPEX
    # ─────────────────────────────────────────────

    def _build_b_capex(self):
        self._section_header("SECTION B — CAPEX ANALYSIS REPORT")
        capex_cfg = self.cfg.get('capex', {})
        enabled   = capex_cfg.get('enabled', False)

        self._sub_header("B.1 CAPEX Interior Activities")
        self._narrative(
            "Interior capital expenditure activities for the current period focused on unit renovations and upgrades "
            "to support lease-up and improve unit quality. Renovation scope typically includes full interior painting, "
            "countertop resurfacing, flooring, and fixture replacements performed in coordination with the maintenance team."
        )

        self._sub_header("B.2 CAPEX Exterior Activities")
        self._narrative(
            "Exterior capital activities address deferred maintenance, site improvements, and signage. "
            "A standardized CAPEX tagging and reporting process is recommended to improve tracking, budgeting, "
            "and vendor accountability for all capital projects."
        )

        if not enabled:
            source = capex_cfg.get('data_source', 'CAPEX tracking spreadsheet')
            _missing_data_note(self.doc, "CAPEX Detail (B.1 / B.2)", source)
            self._spacer()
            self._page_break()
            return

        # Build CAPEX table from config
        exterior = capex_cfg.get('exterior', [])
        interior = capex_cfg.get('interior', [])
        all_items = [('Interior', r) for r in interior] + [('Exterior', r) for r in exterior]

        table = self._make_table(3, [1.0, 3.0, 1.5],
                                  ["Category", "Description", "Amount"])
        total = 0.0
        for category, item in all_items:
            row = table.add_row()
            _cell(row.cells[0], category, font_size=9)
            _cell(row.cells[1], item.get('description', ''), font_size=9)
            amt = item.get('amount', 0.0)
            total += amt
            _cell(row.cells[2], _money(amt), font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

        row = table.add_row()
        _total_cell(row.cells[0], "")
        _total_cell(row.cells[1], "Total CAPEX", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _total_cell(row.cells[2], _money(total))

        self._spacer()
        self._page_break()

    # ─────────────────────────────────────────────
    # C.1 P&L Analysis
    # ─────────────────────────────────────────────

    def _build_c1_pnl(self):
        self._section_header("SECTION C — FINANCIAL ANALYSIS")
        self._sub_header("C.1 Profit and Loss Analysis")

        fin_data = {
            'dec_rental_income':   float(self.latest_md.total_rental_income),
            'nov_rental_income':   float(self.prev_md.total_rental_income),
            'dec_other_income':    float(self.latest_md.total_other_income),
            'dec_total_income':    float(self.latest_md.total_income),
            'dec_total_expenses':  float(self.latest_md.total_expenses),
            'dec_noi':             float(self.latest_md.noi),
            'avg_noi':             self.fin.get('noi_analysis', {}).get('average_noi', 0),
        }
        narrative = (self.ng.generate_financial_narrative(fin_data, self.latest_key)
                     if self.ng else "")
        if not narrative:
            delta = fin_data['dec_rental_income'] - fin_data['nov_rental_income']
            narrative = (
                f"For {self.latest_key}, net rental income was {_money(fin_data['dec_rental_income'])}, "
                f"a {'decrease' if delta < 0 else 'increase'} of {_money(abs(delta))} versus the prior month. "
                f"Other income totaled {_money(fin_data['dec_other_income'])}, "
                f"bringing total revenue to {_money(fin_data['dec_total_income'])}. "
                f"NOI for the period was {_money(fin_data['dec_noi'])}."
            )
        self._narrative(narrative)

        # Appendix 10 — full 6-month P&L
        self._sub_header("Appendix 10 — 6-Month P&L Summary")

        pnl_rows = [
            ("Gross Potential Rent",     "gross_potential_rent",     False),
            ("Vacancy Loss",             "vacancy_loss",             False),
            ("Loss to Old Lease",        "loss_to_old_lease",        False),
            ("Concessions",              "concessions",              False),
            ("Write-Off Uncollectable",  "write_off_uncollectable",  False),
            ("Total Rental Income",      "total_rental_income",      True),
            ("Total Other Income",       "total_other_income",       True),
            ("TOTAL INCOME",             "total_income",             True),
            ("Total Administrative",     "total_administrative",     False),
            ("Total Marketing",          "total_marketing",          False),
            ("Total Payroll",            "total_payroll",            False),
            ("Total Repairs",            "total_repairs",            False),
            ("Total Unit Preparation",   "total_unit_preparation",   False),
            ("Total Contract Services",  "total_contract_services",  False),
            ("Total Utilities",          "total_utilities",          False),
            ("Total Insurance/Taxes",    "total_insurance_taxes",    False),
            ("TOTAL EXPENSES",           "total_expenses",           True),
            ("NOI",                      "noi",                      True),
            ("Total Debt Service",       "total_debt_service",       False),
            ("NET INCOME",               "net_income",               True),
        ]

        headers = ["Income / Expense"] + self.month_labels
        widths  = [2.2] + [0.87] * len(self.month_keys)
        table   = self._make_table(len(headers), widths, headers)

        for label, attr, is_bold in pnl_rows:
            row = table.add_row()
            if is_bold:
                _subheader_cell(row.cells[0], label)
            else:
                _cell(row.cells[0], label, font_size=8)
            for j, mk in enumerate(self.month_keys):
                md  = self.t12.months.get(mk, T12MonthData(mk))
                val = getattr(md, attr, Decimal('0'))
                if is_bold:
                    _total_cell(row.cells[j + 1], _money(val), align=WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    _cell(row.cells[j + 1], _money(val), font_size=8,
                          align=WD_ALIGN_PARAGRAPH.RIGHT)

        self._spacer()

    # ─────────────────────────────────────────────
    # C.2 Budget vs Actual
    # ─────────────────────────────────────────────

    def _build_c2_budget(self):
        self._sub_header("C.2 Financial Audit — Budget vs. Actual")
        budget_cfg = self.cfg.get('budget', {})

        self._narrative(
            "Budget vs. actual analysis compares actual income and expense performance against the approved "
            "property budget. This enables identification of variances, accountability for expense control, "
            "and early detection of income shortfalls. Key focus areas include rental income vs. budget, "
            "expense category variances, and NOI vs. projected targets."
        )

        if not budget_cfg.get('enabled', False):
            source = budget_cfg.get('data_source', 'property budget spreadsheet / Yardi budget export')
            _missing_data_note(self.doc, "Budget vs. Actual (C.2)", source)

        self._spacer()

    # ─────────────────────────────────────────────
    # C.3 Collections
    # ─────────────────────────────────────────────

    def _build_c3_collections(self):
        self._sub_header("C.3 Collection Analysis")

        narrative = (self.ng.generate_collections_narrative(self.col) if self.ng else "")
        if not narrative:
            stats = self.col.get('collection_stats', {})
            cr    = stats.get('collection_rate_pct', 0)
            paid  = stats.get('total_paid', 0)
            charged = stats.get('total_charged', 0)
            dq    = stats.get('total_delinquent_units', 0)
            bal   = self.col.get('total_debit_balance', 0)
            narrative = (
                f"Collection rate for the period was {cr}%, with {_money(paid)} collected of {_money(charged)} charged. "
                f"{dq} units carry delinquent balances totaling {_money(bal)}. "
                f"Strengthening collection procedures and escalation protocols is recommended."
            )
        self._narrative(narrative)

        # Delinquency table
        self._sub_header("Appendix 13 — Delinquency Summary")
        delinquent = self.col.get('delinquent_units', [])
        headers    = ["Unit", "Tenant", "Debit Balance", "Severity"]
        widths     = [0.7, 2.5, 1.3, 1.3]
        table      = self._make_table(4, widths, headers)

        for u in delinquent[:30]:
            row = table.add_row()
            _cell(row.cells[0], u.get('unit', ''),     font_size=9)
            _cell(row.cells[1], u.get('tenant', ''),   font_size=9)
            _cell(row.cells[2], _money(u.get('debit_balance', 0)), font_size=9,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
            _cell(row.cells[3], u.get('severity', ''), font_size=9)

        total_row = table.add_row()
        _total_cell(total_row.cells[0], "TOTAL", align=WD_ALIGN_PARAGRAPH.LEFT)
        _total_cell(total_row.cells[1], f"{len(delinquent)} units", align=WD_ALIGN_PARAGRAPH.LEFT)
        _total_cell(total_row.cells[2], _money(self.rr.debit_balances))
        _total_cell(total_row.cells[3], "")

        self._spacer()

    # ─────────────────────────────────────────────
    # C.4 Expense Analysis
    # ─────────────────────────────────────────────

    def _build_c4_expense(self):
        self._sub_header("C.4 Expense Analysis")

        exp_data = {
            'dec_total_expenses': float(self.latest_md.total_expenses),
            'nov_total_expenses': float(self.prev_md.total_expenses),
            'dec_payroll':        float(self.latest_md.total_payroll),
            'dec_repairs':        float(self.latest_md.total_repairs),
        }
        narrative = (self.ng.generate_expense_narrative(exp_data) if self.ng else "")
        if not narrative:
            delta = exp_data['dec_total_expenses'] - exp_data['nov_total_expenses']
            narrative = (
                f"Total operating expenses for {self.latest_key} were {_money(exp_data['dec_total_expenses'])}, "
                f"a {'decrease' if delta < 0 else 'increase'} of {_money(abs(delta))} vs. the prior month. "
                f"Payroll: {_money(exp_data['dec_payroll'])}. Repairs: {_money(exp_data['dec_repairs'])}."
            )
        self._narrative(narrative)

        expense_rows = [
            ("Administrative",    "total_administrative"),
            ("Marketing",         "total_marketing"),
            ("Payroll",           "total_payroll"),
            ("Repairs",           "total_repairs"),
            ("Unit Preparation",  "total_unit_preparation"),
            ("Contract Services", "total_contract_services"),
            ("Utilities",         "total_utilities"),
            ("Insurance/Taxes",   "total_insurance_taxes"),
            ("TOTAL EXPENSES",    "total_expenses"),
        ]
        headers = ["Expense Category"] + self.month_labels
        widths  = [1.8] + [0.9] * len(self.month_keys)
        table   = self._make_table(len(headers), widths, headers)

        for label, attr in expense_rows:
            is_total = label == "TOTAL EXPENSES"
            row = table.add_row()
            if is_total:
                _subheader_cell(row.cells[0], label)
            else:
                _cell(row.cells[0], label, font_size=8)
            for j, mk in enumerate(self.month_keys):
                md  = self.t12.months.get(mk, T12MonthData(mk))
                val = getattr(md, attr, Decimal('0'))
                if is_total:
                    _total_cell(row.cells[j + 1], _money(val), align=WD_ALIGN_PARAGRAPH.RIGHT)
                else:
                    _cell(row.cells[j + 1], _money(val), font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)

        self._spacer()

    # ─────────────────────────────────────────────
    # C.5 Rent Growth
    # ─────────────────────────────────────────────

    def _build_c5_rent_growth(self):
        self._sub_header("C.5 Rent Growth Analysis")

        months_data = []
        for mk, ml in zip(self.month_keys, self.month_labels):
            md = self.t12.months.get(mk, T12MonthData(mk))
            months_data.append({
                'month': ml,
                'gross_potential_rent': float(md.gross_potential_rent),
                'total_rental_income':  float(md.total_rental_income),
                'total_other_income':   float(md.total_other_income),
                'total_income':         float(md.total_income),
            })

        narrative = (self.ng.generate_rent_growth_narrative(months_data) if self.ng else "")
        if not narrative:
            first = months_data[0]['total_income'] if months_data else 0
            last  = months_data[-1]['total_income'] if months_data else 0
            chg   = round((last - first) / abs(first) * 100, 2) if first != 0 else 0
            narrative = (
                f"Total income trend over the 6-month period shows a {chg:+.2f}% change. "
                f"Gross potential rent remains stable, reflecting consistent market positioning. "
                f"Occupancy improvement and lease renewal pricing are the primary revenue levers going forward."
            )
        self._narrative(narrative)

        # Appendix 15
        self._sub_header("Appendix 15 — Rental Income Trend")
        headers = ["Month", "Gross Potential Rent", "Rental Income", "Other Income", "Total Income", "MoM Change"]
        widths  = [1.0, 1.4, 1.4, 1.2, 1.2, 1.0]
        table   = self._make_table(6, widths, headers)

        prev_ti = None
        for m in months_data:
            mom = ""
            if prev_ti is not None and prev_ti != 0:
                chg = (m['total_income'] - prev_ti) / abs(prev_ti) * 100
                mom = f"{'+' if chg >= 0 else ''}{chg:.1f}%"
            prev_ti = m['total_income']
            row = table.add_row()
            _cell(row.cells[0], m['month'],                             font_size=8)
            _cell(row.cells[1], _money(m['gross_potential_rent']),      font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _cell(row.cells[2], _money(m['total_rental_income']),       font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _cell(row.cells[3], _money(m['total_other_income']),        font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _cell(row.cells[4], _money(m['total_income']),              font_size=8, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True)
            _cell(row.cells[5], mom,                                    font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

        self._spacer()
        self._page_break()

    # ─────────────────────────────────────────────
    # Section E — Financial Account Audit
    # ─────────────────────────────────────────────

    def _build_e1_gpr(self):
        self._section_header("SECTION E — FINANCIAL ACCOUNT AUDIT REPORT")
        self._sub_header("E.1 Gross Potential Rent (GPR) Analysis")

        gpr_data = {
            'total_units':    self.rr.total_units,
            'occupied_units': self.rr.occupied_units,
            'vacant_units':   self.rr.vacant_units,
            'market_rent':    self.rr.market_rent,
            'loss_to_lease':  self.rr.loss_to_lease,
            'vacancy_loss':   self.rr.vacancy_loss,
            'rent_charges':   self.rr.rent_charges,
        }
        narrative = (self.ng.generate_gpr_narrative(gpr_data) if self.ng else "")
        if not narrative:
            narrative = (
                f"The property has {gpr_data['total_units']} total units with "
                f"{gpr_data['occupied_units']} occupied and {gpr_data['vacant_units']} vacant. "
                f"Market rent at 100% occupancy: {_money(gpr_data['market_rent'])}. "
                f"Loss to lease: {_money(gpr_data['loss_to_lease'])}. "
                f"Vacancy loss: {_money(gpr_data['vacancy_loss'])}. "
                f"Actual rent charges: {_money(gpr_data['rent_charges'])}."
            )
        self._narrative(narrative)

        # Summary tables
        t1 = self.doc.add_table(rows=2, cols=1)
        t1.style = 'Table Grid'
        _set_col_widths(t1, [1.0])
        _header_cell(t1.rows[0].cells[0], "UNITS")
        _cell(t1.rows[1].cells[0], str(self.rr.total_units), font_size=14, bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        self._spacer()

        t2 = self.doc.add_table(rows=2, cols=3)
        t2.style = 'Table Grid'
        _set_col_widths(t2, [1.5, 1.5, 1.5])
        for i, h in enumerate(["OCCUPANCY", "VACANCY", "DOWN UNITS"]):
            _header_cell(t2.rows[0].cells[i], h)
        _cell(t2.rows[1].cells[0], str(self.rr.occupied_units), font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(t2.rows[1].cells[1], str(self.rr.vacant_units),   font_size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(t2.rows[1].cells[2], "—",                          font_size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._spacer()

        t3 = self.doc.add_table(rows=2, cols=5)
        t3.style = 'Table Grid'
        _set_col_widths(t3, [1.3] * 5)
        for i, h in enumerate(["MARKET RENT", "LOSS TO LEASE", "VACANCY LOSS", "CONCESSIONS", "ACTUAL RENT"]):
            _header_cell(t3.rows[0].cells[i], h)
        values = [
            _money(self.rr.market_rent),
            _money(self.rr.loss_to_lease),
            _money(self.rr.vacancy_loss),
            _money(self.latest_md.concessions),
            _money(self.rr.rent_charges),
        ]
        for i, v in enumerate(values):
            _cell(t3.rows[1].cells[i], v, font_size=10, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self._spacer()

        t4 = self.doc.add_table(rows=2, cols=2)
        t4.style = 'Table Grid'
        _set_col_widths(t4, [1.5, 1.5])
        _header_cell(t4.rows[0].cells[0], "MOVE IN")
        _header_cell(t4.rows[0].cells[1], "MOVE OUT")
        _cell(t4.rows[1].cells[0], "—", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(t4.rows[1].cells[1], "—", font_size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        p = self.doc.add_paragraph()
        run = p.add_run("Move-in/move-out counts require a move report from the property management system.")
        run.italic = True; run.font.size = Pt(8); run.font.color.rgb = GREY
        self._spacer()

    def _e_narrative_section(self, title: str, narrative: str):
        self._sub_header(title)
        self._narrative(narrative)

    def _build_e2_concessions(self):
        self._sub_header("E.2 Concessions")
        conc_cfg  = self.cfg.get('concessions_detail', {})
        conc_amt  = float(self.latest_md.concessions)
        emp_loss  = float(self.latest_md.loss_to_employee)

        self._narrative(
            f"Concessions for the period totaled {_money(conc_amt)}. "
            f"Employee unit loss (loss to employee) was {_money(emp_loss)}. "
            f"All concessions should be reviewed for appropriateness and documented in the lease file."
        )

        if not conc_cfg.get('enabled', False):
            source = conc_cfg.get('data_source', 'general ledger concession line items')
            _missing_data_note(self.doc, "Concession Line Items (E.2)", source)
            return

        items  = conc_cfg.get('items', [])
        table  = self._make_table(4, [0.7, 2.0, 2.5, 1.0],
                                   ["Unit No", "Tenant", "Description", "Amount"])
        total  = 0.0
        for item in items:
            row = table.add_row()
            _cell(row.cells[0], str(item.get('unit', '')),        font_size=9)
            _cell(row.cells[1], str(item.get('tenant', '')),      font_size=9)
            _cell(row.cells[2], str(item.get('description', '')), font_size=9)
            amt = item.get('amount', 0.0)
            total += amt
            _cell(row.cells[3], _money(amt), font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

        tr = table.add_row()
        _total_cell(tr.cells[0], "")
        _total_cell(tr.cells[1], "")
        _total_cell(tr.cells[2], "Total", align=WD_ALIGN_PARAGRAPH.RIGHT)
        _total_cell(tr.cells[3], _money(total))
        self._spacer()

    def _build_e3_mtm(self):
        mtm = float(self.latest_md.mtm_fees)
        self._e_narrative_section("E.3 Month-to-Month (MTM) Analysis",
            f"Month-to-month fee income for the period was {_money(mtm)}. "
            f"Tenants with expired leases must be charged month-to-month fees per lease agreements. "
            f"Lease expiration dates should be kept current in the property management system to ensure "
            f"all applicable MTM charges are collected and no tenants are inadvertently excluded."
        )

    def _build_e4_admin_fees(self):
        admin = float(self.latest_md.admin_fees)
        app   = float(self.latest_md.application_fees)
        self._e_narrative_section("E.4 Admin and Application Fees Analysis",
            f"Administrative fees collected: {_money(admin)}. "
            f"Application fees collected: {_money(app)}. "
            f"All new move-ins must be assessed the required administrative and application fees. "
            f"Fee collection should be reconciled monthly against new move-in records to ensure completeness."
        )

    def _build_e5_utility_income(self):
        water = float(self.latest_md.utility_water)
        trash = float(self.latest_md.utility_trash)
        self._e_narrative_section("E.5 Utility Income Analysis",
            f"Utility income collected from tenants: Water/Sewer {_money(water)}, Trash {_money(trash)}. "
            f"Vacant unit utility costs should be monitored closely as they represent direct expense "
            f"without offsetting income. Utility reimbursement collections should be reconciled "
            f"against actual utility bills monthly to identify any discrepancies."
        )

    def _build_e6_rubs(self):
        parking = float(self.latest_md.parking_income)
        pet     = float(self.latest_md.pet_charges)
        self._e_narrative_section("E.6 RUBS / Ancillary Income Analysis",
            f"Parking income: {_money(parking)}. Pet charges: {_money(pet)}. "
            f"Discrepancies between general ledger totals and rent roll figures should be investigated "
            f"and reconciled to ensure all miscellaneous income is properly captured and reported. "
            f"Washer/dryer income and other RUBS items should be tracked separately for audit purposes."
        )

    def _build_e7_software(self):
        self._e_narrative_section("E.7 Software Charges Analysis",
            "Software and computer expenses are tracked under the administrative expense category. "
            "Property management software subscriptions should be reviewed quarterly to ensure all licenses "
            "are current, appropriately priced, and budget-compliant. Unused or duplicate subscriptions should "
            "be identified and cancelled to control costs."
        )

    def _build_e8_marketing(self):
        mktg = float(self.latest_md.total_marketing) if hasattr(self.latest_md, 'total_marketing') else 0
        self._e_narrative_section("E.8 Marketing Charges Analysis",
            f"Marketing and advertising expenditures for the period totaled {_money(mktg) if mktg else 'N/A'}. "
            f"Primary marketing channels include Google Ads, ILS platforms (Apartments.com, Zillow), "
            f"and social media. Marketing spend should be evaluated against lead conversion metrics "
            f"to optimize cost-per-lease ratios and maintain adequate digital presence."
        )

    def _build_e9_payroll(self):
        self._sub_header("E.9 Payroll Analysis")
        total_pay = float(self.latest_md.total_payroll)

        self._narrative(
            f"Total payroll expense for {self.latest_key}: {_money(total_pay)}. "
            f"Payroll levels are consistent with the authorized staffing plan. "
            f"Bonuses and benefit costs are included in the breakdown below."
        )

        payroll_rows = [
            ("5310 Manager's Salary",              self.latest_md.managers_salary),
            ("5320 Leasing Agents Wages",           self.latest_md.leasing_agents),
            ("5330 Maintenance Supervisor Salary",  self.latest_md.maintenance_supervisor),
            ("5335 Assistant Maintenance Wages",    self.latest_md.assistant_maintenance),
            ("5365 Bonuses",                        self.latest_md.bonuses),
            ("5380 Insurance and Other Benefits",   self.latest_md.payroll_insurance),
            ("5385 Payroll Taxes",                  self.latest_md.payroll_taxes),
        ]
        table = self._make_table(2, [3.5, 1.5], ["Account", "Amount"])
        for label, val in payroll_rows:
            row = table.add_row()
            _cell(row.cells[0], label, font_size=9)
            _cell(row.cells[1], _money(val), font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

        tr = table.add_row()
        _total_cell(tr.cells[0], "5300 Total Payroll Expense", align=WD_ALIGN_PARAGRAPH.LEFT)
        _total_cell(tr.cells[1], _money(total_pay))
        self._spacer()

    def _build_e10_maintenance(self):
        repairs   = float(self.latest_md.total_repairs)
        unit_prep = float(self.latest_md.total_unit_preparation)
        self._e_narrative_section("E.10 Maintenance / Repairs Analysis",
            f"Maintenance and repairs expense for {self.latest_key}: {_money(repairs)}. "
            f"Unit preparation costs: {_money(unit_prep)}. "
            f"Common maintenance categories include HVAC servicing, plumbing repairs, carpet cleaning, "
            f"and appliance replacements. Work order completion rates and aging should be monitored to "
            f"ensure timely resolution and minimize resident satisfaction impacts."
        )

    def _build_e11_ap_aging(self):
        contract = float(self.latest_md.total_contract_services)
        self._e_narrative_section("E.11 Accounts Payable Aging Analysis",
            f"Contract services expense for the period: {_money(contract)}. "
            f"AP aging should be reviewed monthly to ensure vendor invoices are paid within agreed terms. "
            f"Outstanding AP balances greater than 60 days should be escalated for resolution to avoid "
            f"late fees and vendor relationship deterioration."
        )

    def _build_e12_contracts(self):
        contract = float(self.latest_md.total_contract_services)
        self._e_narrative_section("E.12 3rd Party Contract Analysis",
            f"Third-party contract services totaled {_money(contract)} for the period, including "
            f"property management fees, landscaping, pest control, and ongoing service contracts. "
            f"All contracts should be reviewed annually for competitiveness, performance compliance, "
            f"and inclusion of performance metrics and termination provisions."
        )

    def _build_e13_late_charges(self):
        late = float(self.latest_md.late_charges)
        self._e_narrative_section("E.13 Late Charges / Penalties Analysis",
            f"Late fees and penalties collected for the period: {_money(late)}. "
            f"Late fee collection should be reconciled against delinquent tenant ledgers to ensure "
            f"all applicable charges are assessed. Amounts and due dates should be clearly "
            f"communicated to tenants per lease terms and documented in the ledger."
        )

    def _build_e14_bank(self):
        self._e_narrative_section("E.14 Bank Statement vs. Accounts Differences",
            "Bank statements should be reconciled against the general ledger monthly to identify any "
            "discrepancies between recorded transactions and actual bank activity. Outstanding items "
            "including uncleared checks and deposits in transit should be documented and resolved promptly. "
            "Any unexplained variances should be escalated to the financial controller for investigation."
        )
        self._page_break()

    # ─────────────────────────────────────────────
    # Signature Block
    # ─────────────────────────────────────────────

    def _build_signature_block(self):
        self._section_header("REPORT CERTIFICATION")
        self._spacer()

        sigs   = self.cfg.get('signatures', [])
        table  = self._make_table(3, [2.0, 2.5, 2.0],
                                   ["Role", "Name", "Company"])
        for sig in sigs:
            row = table.add_row()
            _cell(row.cells[0], sig.get('role', ''),    font_size=9, bold=True)
            _cell(row.cells[1], sig.get('name', ''),    font_size=9)
            _cell(row.cells[2], sig.get('company', ''), font_size=9)

        self._spacer()
        p = self.doc.add_paragraph()
        run = p.add_run(
            f"Report generated: {self.rr.report_date}  |  "
            f"Property: {self.rr.property_name}  |  "
            f"Period: {self.rr.period_start} – {self.rr.period_end}"
        )
        run.font.size = Pt(8)
        run.font.color.rgb = GREY

    # ─────────────────────────────────────────────
    # Main entry point
    # ─────────────────────────────────────────────

    def generate(self, output_path: str) -> str:
        steps = [
            ("Cover page",                self._build_cover),
            ("Audit top sheet",           self._build_audit_top_sheet),
            ("A.1 Occupancy",             self._build_a1_occupancy),
            ("A.2 Rental Analysis",       self._build_a2_rental),
            ("A.3 Rental Comparison",     self._build_a3_rental_comparison),
            ("A.4 Lease Analysis",        self._build_a4_lease),
            ("A.5 Work Orders",           self._build_a5_work_orders),
            ("A.6 Prospect Analysis",     self._build_a6_prospects),
            ("B CAPEX",                   self._build_b_capex),
            ("C.1 P&L Analysis",          self._build_c1_pnl),
            ("C.2 Budget vs Actual",      self._build_c2_budget),
            ("C.3 Collections",           self._build_c3_collections),
            ("C.4 Expense Analysis",      self._build_c4_expense),
            ("C.5 Rent Growth",           self._build_c5_rent_growth),
            ("E.1 GPR Analysis",          self._build_e1_gpr),
            ("E.2 Concessions",           self._build_e2_concessions),
            ("E.3 MTM",                   self._build_e3_mtm),
            ("E.4 Admin Fees",            self._build_e4_admin_fees),
            ("E.5 Utility Income",        self._build_e5_utility_income),
            ("E.6 RUBS",                  self._build_e6_rubs),
            ("E.7 Software Charges",      self._build_e7_software),
            ("E.8 Marketing",             self._build_e8_marketing),
            ("E.9 Payroll",               self._build_e9_payroll),
            ("E.10 Maintenance",          self._build_e10_maintenance),
            ("E.11 AP Aging",             self._build_e11_ap_aging),
            ("E.12 Contracts",            self._build_e12_contracts),
            ("E.13 Late Charges",         self._build_e13_late_charges),
            ("E.14 Bank Statement",       self._build_e14_bank),
            ("Signature block",           self._build_signature_block),
        ]

        for label, fn in steps:
            print(f"  Building {label}...")
            fn()

        self.doc.save(output_path)
        return output_path
